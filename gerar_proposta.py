#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gerador da Proposta Comercial PILAR.

Lê os dados via stdin (JSON), preenche o modelo proposta_modelo.docx e escreve
o .docx resultante em stdout (binário). Todo diagnóstico vai para stderr.

Entrada (stdin), todos os campos opcionais exceto os itens:
{
  "numero_pil": "PIL-016-2026",
  "cliente": "GILIDIVAN",
  "itens": [                          # preço de VENDA (USD); nunca FOB/custo
    {"codigo": "E01", "produto": "TEC...80GSM DES A",
     "quantidade": 30000, "unidade": "M", "pv_unit_usd": 0.81}
  ],
  "grupos": [                         # opcional: grupos do corpo (nome editável
    {"produto": "TEC...80GSM", "qtd": 68000,   #  no painel). Ausente -> agrupa
     "unidade": "M", "pv_unit_usd": 0.81}      #  aqui (agrupar/nome_grupo).
  ],
  "modalidade_frete": "o container de 40hc",   # ou "LCL"
  "pct_sinal": 10,
  "cambio_sinal": 5.1680,
  "data_sinal": "2026-07-14",         # ISO (yyyy-mm-dd)
  "data_venc_sinal": "",              # ISO; vazio -> some "até o dia"
  "cambio_ref": 5.1680, "data_ref": "2026-07-14",
  "dias_antes_desembarque": 20,
  "frete_usd": 150, "prazo_entrega": 60,
  "observacoes": ""
}

Estrutura do .docx:
  - Corpo: um parágrafo por GRUPO (PV unitário em negrito), total emendado no
    último; o Sinal (i) é um único parágrafo contínuo.
  - Anexo (última página): tabela com TODOS os itens + linha TOTAL GERAL.
Cálculos: pv_total_usd = soma(qtd*pv_unit); valor_sinal_usd = pv_total*pct/100;
valor_sinal_brl = valor_sinal_usd * cambio_sinal; pct_saldo = 100 - pct_sinal.
Os campos *_EXTENSO são gerados por extenso em português (sem dependências).
"""
import io
import os
import re
import sys
import json
import copy

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

MODELO = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      'proposta_modelo.docx')

# ── Número por extenso (PT-BR) ───────────────────────────────────────────────
_UNID = ['zero', 'um', 'dois', 'três', 'quatro', 'cinco', 'seis', 'sete',
         'oito', 'nove', 'dez', 'onze', 'doze', 'treze', 'quatorze', 'quinze',
         'dezesseis', 'dezessete', 'dezoito', 'dezenove']
_DEZ = ['', '', 'vinte', 'trinta', 'quarenta', 'cinquenta', 'sessenta',
        'setenta', 'oitenta', 'noventa']
_CENT = ['', 'cento', 'duzentos', 'trezentos', 'quatrocentos', 'quinhentos',
         'seiscentos', 'setecentos', 'oitocentos', 'novecentos']
_ESCALA_S = ['', 'mil', 'milhão', 'bilhão', 'trilhão']
_ESCALA_P = ['', 'mil', 'milhões', 'bilhões', 'trilhões']


def _ate999(n):
    if n == 0:
        return ''
    if n == 100:
        return 'cem'
    partes = []
    c = n // 100
    resto = n % 100
    if c:
        partes.append(_CENT[c])
    if resto:
        if resto < 20:
            partes.append(_UNID[resto])
        else:
            d, u = resto // 10, resto % 10
            partes.append(_DEZ[d] + (' e ' + _UNID[u] if u else ''))
    return ' e '.join(partes)


def extenso_int(n):
    n = int(round(n))
    if n == 0:
        return 'zero'
    if n < 0:
        return 'menos ' + extenso_int(-n)
    grupos = []
    while n > 0:
        grupos.append(n % 1000)
        n //= 1000
    termos = []
    for i in range(len(grupos) - 1, -1, -1):
        g = grupos[i]
        if g == 0:
            continue
        if i == 0:
            termos.append(_ate999(g))
        elif i == 1:
            termos.append('mil' if g == 1 else _ate999(g) + ' mil')
        else:
            nome = _ESCALA_S[i] if g == 1 else _ESCALA_P[i]
            termos.append(_ate999(g) + ' ' + nome)
    if len(termos) == 1:
        return termos[0]
    # valor do grupo de menor ordem presente define o conector final
    ultimo = next(g for g in grupos if g)
    conector = ' e ' if (ultimo < 100 or ultimo % 100 == 0) else ' '
    return ', '.join(termos[:-1]) + conector + termos[-1]


def _cap(s):
    return s[0].upper() + s[1:] if s else s


def _parte_centavos(valor):
    inteiro = int(valor)
    cent = int(round((valor - inteiro) * 100))
    if cent == 100:
        inteiro += 1
        cent = 0
    return inteiro, cent


def extenso_moeda(valor, sing, plur, csing, cplur):
    valor = round(float(valor) + 1e-9, 2)
    inteiro, cent = _parte_centavos(valor)
    partes = []
    if inteiro:
        partes.append(extenso_int(inteiro) + ' ' + (sing if inteiro == 1 else plur))
    if cent:
        partes.append(extenso_int(cent) + ' ' + (csing if cent == 1 else cplur))
    if not partes:
        partes.append('zero ' + plur)
    return _cap(' e '.join(partes))


def extenso_brl(v):
    return extenso_moeda(v, 'real', 'reais', 'centavo', 'centavos')


def extenso_usd(v):
    return extenso_moeda(v, 'dólar', 'dólares',
                         'centavo de dólar', 'centavos de dólar')


# ── Formatação numérica pt-BR ────────────────────────────────────────────────
def fmt_num(v, dec=2):
    s = ('{:,.%df}' % dec).format(float(v))
    return s.replace(',', '#').replace('.', ',').replace('#', '.')


def fmt_preco(v):
    """Preço unitário: 2 casas quando >= USD 1 (5.733,62); 4 casas quando
    < USD 1 (0,9660)."""
    v = float(v)
    return fmt_num(v, 2 if abs(v) >= 1 else 4)


def fmt_qtd(v):
    """Quantidade: inteiro quando exata; senão 2 casas."""
    v = float(v)
    if abs(v - round(v)) < 1e-9:
        return fmt_num(v, 0).split(',')[0]
    return fmt_num(v, 2)


_MESES = ['', 'Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho',
          'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro']


def fmt_data_curta(iso):
    m = re.match(r'(\d{4})-(\d{2})-(\d{2})', str(iso or ''))
    if not m:
        return str(iso or '')
    y, mo, d = m.groups()
    return '%s/%s/%s' % (d, mo, y)


def fmt_data_longa(iso):
    m = re.match(r'(\d{4})-(\d{2})-(\d{2})', str(iso or ''))
    if not m:
        return str(iso or '')
    y, mo, d = m.groups()
    return '%d de %s de %s' % (int(d), _MESES[int(mo)], y)


def num(v, default=0.0):
    try:
        if v is None or v == '':
            return default
        return float(str(v).replace('.', '').replace(',', '.')) \
            if isinstance(v, str) and ',' in str(v) else float(v)
    except Exception:
        return default


# Normaliza a lista de itens da proposta (preço de venda em USD; nunca FOB/custo).
def itens_norm(d):
    out = []
    for it in (d.get('itens') or []):
        it = it or {}
        prod = str(it.get('produto', '') or '').strip()
        cod = str(it.get('codigo', '') or '').strip()
        qtd = num(it.get('quantidade'), 0)
        unidade = str(it.get('unidade', '') or '').strip()
        pvu = num(it.get('pv_unit_usd'), 0)
        if not prod and qtd <= 0:
            continue
        out.append({'codigo': cod, 'produto': prod, 'qtd': qtd,
                    'unidade': unidade, 'pv_unit_usd': pvu,
                    'total_usd': round(qtd * pvu, 2)})
    return out


def singular_un(u):
    u = (u or '').strip()
    return u[:-1] if len(u) > 1 and u.lower().endswith('s') else u


# ── Agrupamento (corpo da proposta) ──────────────────────────────────────────
# Nome do grupo = descrição truncada logo após o token "...GSM"; sem match, a
# descrição inteira. Grupo = mesmo nome + mesma unidade + mesmo PV unitário.
_GSM_RE = re.compile(r'^(.*?\d+\s*GSM)', re.IGNORECASE)


def nome_grupo(produto):
    m = _GSM_RE.match(produto or '')
    return (m.group(1).strip() if m else (produto or '').strip())


def agrupar(itens):
    ordem, grupos = [], {}
    for it in itens:
        nome = nome_grupo(it['produto'])
        key = (nome, it['unidade'], round(it['pv_unit_usd'], 4))
        if key not in grupos:
            grupos[key] = {'produto': nome, 'qtd': 0.0,
                           'unidade': it['unidade'],
                           'pv_unit_usd': it['pv_unit_usd']}
            ordem.append(key)
        grupos[key]['qtd'] += it['qtd']
    return [grupos[k] for k in ordem]


# Usa os grupos vindos do painel (nome editável) quando presentes; senão agrupa.
def grupos_norm(d, itens):
    raw = d.get('grupos')
    if isinstance(raw, list) and raw:
        out = []
        for g in (raw or []):
            g = g or {}
            prod = str(g.get('produto', '') or '').strip()
            qtd = num(g.get('qtd'), 0)
            if not prod and qtd <= 0:
                continue
            out.append({'produto': prod, 'qtd': qtd,
                        'unidade': str(g.get('unidade', '') or '').strip(),
                        'pv_unit_usd': num(g.get('pv_unit_usd'), 0)})
        if out:
            return out
    return agrupar(itens)


# Segmentos (texto, negrito) de um parágrafo de grupo. PV unit em negrito.
# No último grupo, emenda a frase do total após vírgula.
def linha_grupo_segs(g, ultimo, mapa):
    qtd_txt = fmt_qtd(g['qtd'])
    pvu = fmt_preco(g['pv_unit_usd'])
    ext = extenso_usd(g['pv_unit_usd'])
    un = g['unidade']
    if un:
        segs = [('%s %s de %s ' % (qtd_txt, un, g['produto']), False),
                ('USD %s' % pvu, True),
                (' por %s (%s)' % (singular_un(un), ext), False)]
    else:
        segs = [('%s de %s ' % (qtd_txt, g['produto']), False),
                ('USD %s' % pvu, True),
                (' (%s)' % ext, False)]
    if ultimo:
        segs.append((', O valor total estimado do pedido antes do embarque é '
                     'de USD %s (%s).' % (mapa['{{PV_TOTAL_USD}}'],
                                          mapa['{{PV_TOTAL_USD_EXTENSO}}']),
                     False))
    return segs


def montar_mapa(d, itens):
    cliente = str(d.get('cliente', '') or '')
    numero_pil = str(d.get('numero_pil', '') or '')
    qtd_containers = num(d.get('qtd_containers'), 1)
    tipo_container = str(d.get('tipo_container', '40HC') or '40HC')
    pct_sinal = num(d.get('pct_sinal'), 20)
    pct_saldo = 100 - pct_sinal
    cambio = num(d.get('cambio_sinal'), 0)
    cambio_ref = num(d.get('cambio_ref'), 0)
    dias_antes = num(d.get('dias_antes_desembarque'), 20)
    frete = num(d.get('frete_usd'), 0)
    prazo = num(d.get('prazo_entrega'), 60)
    modalidade = (str(d.get('modalidade_frete', '') or '').strip()
                  or 'o container de 40hc')
    obs = str(d.get('observacoes', '') or '').strip()

    # Preço de venda em USD (âncora). R$ é apenas referência do dia.
    pv_total_usd = round(sum(i['total_usd'] for i in itens), 2)
    valor_sinal_usd = round(pv_total_usd * pct_sinal / 100.0, 2)
    valor_sinal_brl = round(valor_sinal_usd * cambio, 2)
    pv_total_brl_ref = round(pv_total_usd * cambio_ref, 2)

    def pct_txt(p):
        return '%s%% (%s por cento)' % (fmt_num(p, 0).split(',')[0],
                                        extenso_int(p))

    def dias_txt(n):
        return '%s (%s)' % (fmt_num(n, 0).split(',')[0], extenso_int(n))

    return {
        '{{CLIENTE}}': cliente,
        '{{NUMERO_PIL}}': numero_pil,
        '{{QTD_CONTAINERS}}': fmt_num(qtd_containers, 0).split(',')[0],
        '{{TIPO_CONTAINER}}': tipo_container,
        '{{PV_TOTAL_USD}}': fmt_num(pv_total_usd, 2),
        '{{PV_TOTAL_USD_EXTENSO}}': extenso_usd(pv_total_usd),
        '{{PV_TOTAL_BRL_REF}}': fmt_num(pv_total_brl_ref, 2),
        '{{CAMBIO_REF}}': fmt_num(cambio_ref, 4),
        '{{DATA_REF}}': fmt_data_curta(d.get('data_ref')),
        '{{PCT_SINAL}}': pct_txt(pct_sinal),
        '{{PCT_SALDO}}': pct_txt(pct_saldo),
        '{{VALOR_SINAL_USD}}': fmt_num(valor_sinal_usd, 2),
        '{{VALOR_SINAL_USD_EXTENSO}}': extenso_usd(valor_sinal_usd),
        '{{CAMBIO_SINAL}}': fmt_num(cambio, 4),
        '{{DATA_SINAL}}': fmt_data_curta(d.get('data_sinal')),
        '{{VALOR_SINAL_BRL}}': fmt_num(valor_sinal_brl, 2),
        '{{VALOR_SINAL_BRL_EXTENSO}}': extenso_brl(valor_sinal_brl),
        '{{DATA_VENC_SINAL}}': fmt_data_longa(d.get('data_venc_sinal')),
        '{{DIAS_ANTES_DESEMBARQUE}}': dias_txt(dias_antes),
        '{{FRETE_USD}}': fmt_num(frete, 0).split(',')[0],
        '{{MODALIDADE_FRETE}}': modalidade,
        '{{PRAZO_ENTREGA}}': dias_txt(prazo),
        '{{OBSERVACOES}}': obs,
    }


# Monta um w:r com a formatação-base do corpo (Arial 9 / sz 18) e negrito
# opcional — mesma estrutura (rFonts + b + sz) dos runs do modelo.
def _mk_run(text, bold):
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Arial'); rf.set(qn('w:hAnsi'), 'Arial')
    rpr.append(rf)
    b = OxmlElement('w:b')
    if not bold:
        b.set(qn('w:val'), '0')
    rpr.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t)
    return r


# Expande {{GRUPOS}} clonando o parágrafo-placeholder (recuo/estilo do modelo)
# — uma cópia por grupo, com o PV unitário em negrito.
def expandir_grupos(doc, grupos, mapa):
    for p in list(doc.paragraphs):
        if '{{GRUPOS}}' not in p.text:
            continue
        n = len(grupos)
        for i, g in enumerate(grupos):
            novo = copy.deepcopy(p._p)
            for r in novo.findall(qn('w:r')):
                novo.remove(r)
            for text, bold in linha_grupo_segs(g, i == n - 1, mapa):
                novo.append(_mk_run(text, bold))
            p._p.addprevious(novo)
        p._p.getparent().remove(p._p)
        return


def substituir(doc, mapa):
    venc_vazio = not mapa['{{DATA_VENC_SINAL}}']
    for p in list(doc.paragraphs):
        txt = p.text
        # Parágrafo exclusivo das observações: remove se vazio
        if '{{OBSERVACOES}}' in txt:
            if not mapa['{{OBSERVACOES}}']:
                p._element.getparent().remove(p._element)
                continue
        for run in p.runs:
            novo = run.text
            # Sinal sem vencimento: não deixar "à vista até o dia  na conta".
            if venc_vazio and 'à vista até o dia' in novo:
                novo = novo.replace('à vista até o dia ', 'à vista')
            if '{{' in novo:
                for k, v in mapa.items():
                    if k in novo:
                        novo = novo.replace(k, v)
            if novo != run.text:
                run.text = novo


# ── Anexo (última página): tabela com TODOS os itens ─────────────────────────
def _cell(cell, text, bold=False, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'right': WD_ALIGN_PARAGRAPH.RIGHT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}.get(
                       align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is not None:
        rf.set(qn('w:cs'), 'Arial')
        rf.set(qn('w:eastAsia'), 'Arial')


def _tabela_bordas(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'auto')
        borders.append(e)
    tblPr.append(borders)


def adicionar_anexo(doc, itens):
    if not itens:
        return
    mostrar_cod = any(it.get('codigo') for it in itens)

    quebra = doc.add_paragraph()
    quebra.add_run().add_break(WD_BREAK.PAGE)

    titulo = doc.add_paragraph()
    rt = titulo.add_run('Anexo — Relação de Itens')
    rt.bold = True
    rt.font.name = 'Arial'
    rt.font.size = Pt(12)

    cols = (['Código'] if mostrar_cod else []) + \
        ['Descrição', 'Qtd', 'Un', 'PV Unit. USD', 'Total USD']
    table = doc.add_table(rows=1, cols=len(cols))
    table.autofit = True
    _tabela_bordas(table)

    aligns = (['left'] if mostrar_cod else []) + \
        ['left', 'right', 'center', 'right', 'right']
    hdr = table.rows[0].cells
    for j, nome in enumerate(cols):
        _cell(hdr[j], nome, bold=True, align=aligns[j])

    tot_qtd = 0.0
    tot_usd = 0.0
    for it in itens:
        tot_qtd += it['qtd']
        tot_usd += it['total_usd']
        vals = ([it.get('codigo', '')] if mostrar_cod else []) + [
            it['produto'], fmt_qtd(it['qtd']), it['unidade'],
            fmt_preco(it['pv_unit_usd']), fmt_num(it['total_usd'], 2)]
        cells = table.add_row().cells
        for j, val in enumerate(vals):
            _cell(cells[j], val, align=aligns[j])

    # TOTAL GERAL (negrito)
    base = ['TOTAL GERAL'] if not mostrar_cod else ['', 'TOTAL GERAL']
    tot = base + [fmt_qtd(tot_qtd), '', '', fmt_num(round(tot_usd, 2), 2)]
    cells = table.add_row().cells
    for j, val in enumerate(tot):
        _cell(cells[j], val, bold=True, align=aligns[j])


def main():
    raw = sys.stdin.read()
    try:
        d = json.loads(raw) if raw.strip() else {}
    except Exception as e:
        sys.stderr.write('JSON inválido em stdin: %s\n' % e)
        sys.exit(1)
    if not isinstance(d, dict):
        d = {}

    if not os.path.isfile(MODELO):
        sys.stderr.write('Modelo não encontrado: %s\n' % MODELO)
        sys.exit(1)

    itens = itens_norm(d)
    grupos = grupos_norm(d, itens)
    mapa = montar_mapa(d, itens)
    doc = Document(MODELO)
    expandir_grupos(doc, grupos, mapa)
    substituir(doc, mapa)
    adicionar_anexo(doc, itens)

    buf = io.BytesIO()
    doc.save(buf)
    sys.stdout.buffer.write(buf.getvalue())


if __name__ == '__main__':
    main()
