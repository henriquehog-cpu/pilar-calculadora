#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Prepara o modelo da Proposta Comercial.

Abre o .docx original (com os valores reais do PIL-003-2026 / JUMA), substitui
os trechos variáveis por placeholders {{...}} — cada placeholder fica isolado em
um run limpo, de modo que gerar_proposta.py consiga substituí-lo com um simples
run.text.replace() preservando a formatação (negrito) do parágrafo.

Uso:
  python3 _preparar_modelo.py [ORIGEM.docx] [DESTINO.docx]

Defaults:
  ORIGEM  = ~/Downloads/Proposta comercial Pilar - PIL-003-2026 - JUMA.docx
  DESTINO = ./proposta_modelo.docx
"""
import os
import sys
from docx import Document

ORIGEM_PADRAO = os.path.expanduser(
    '~/Downloads/Proposta comercial Pilar - PIL-003-2026 - JUMA.docx')
DESTINO_PADRAO = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                              'proposta_modelo.docx')


def _achar(doc, predicado):
    for p in doc.paragraphs:
        if predicado(p.text):
            return p
    return None


def _set_runs(p, segments):
    """Reescreve os runs do parágrafo preservando fonte (name/size) do 1º run."""
    src = p.runs[0] if p.runs else None
    fname = src.font.name if src else None
    fsize = src.font.size if src else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for text, bold in segments:
        r = p.add_run(text)
        r.bold = bold
        if fname:
            r.font.name = fname
        if fsize:
            r.font.size = fsize


def _inserir_antes(target, segments):
    novo = target.insert_paragraph_before()
    novo.style = target.style
    for text, bold in segments:
        r = novo.add_run(text)
        r.bold = bold
    return novo


def main():
    origem = sys.argv[1] if len(sys.argv) > 1 else ORIGEM_PADRAO
    destino = sys.argv[2] if len(sys.argv) > 2 else DESTINO_PADRAO

    if not os.path.isfile(origem):
        sys.stderr.write('Origem não encontrada: %s\n' % origem)
        sys.exit(1)

    doc = Document(origem)

    # ── Saudação ──────────────────────────────────────────────────────────────
    p = _achar(doc, lambda t: t.strip() == 'Caro Cliente.')
    if p:
        _set_runs(p, [('Caro ', False), ('{{CLIENTE}}', False), ('.', False)])

    # ── Linha de referência (número PIL) antes da saudação ──────────────────────
    p_saud = _achar(doc, lambda t: t.strip().startswith('Caro '))
    if p_saud:
        _inserir_antes(p_saud, [('Ref.: Proposta Comercial ', True),
                                ('{{NUMERO_PIL}}', True)])
        _inserir_antes(p_saud, [('', False)])  # linha em branco de respiro

    # ── Descrição da mercadoria (parágrafo todo em negrito) ─────────────────────
    p = _achar(doc, lambda t: t.strip().startswith('1 Container de 40'))
    if p:
        _set_runs(p, [(
            '1 Container de 40HC contendo {{QTD_TOTAL}} {{UNIDADE}} de '
            '{{DESCRICAO_RESUMIDA}}.', True)])

    # ── Valores finais ──────────────────────────────────────────────────────────
    p = _achar(doc, lambda t: 'Os valores finais para fornecimento' in t)
    if p:
        _set_runs(p, [
            ('Os valores finais para fornecimento do descrito acima, na sede da '
             'sua empresa são de USD {{FOB_UNIT}} ', False),
            ('({{FOB_UNIT_EXTENSO}})', True),
            (' para {{QTD_TOTAL}} {{UNIDADE}} de {{DESCRICAO_RESUMIDA}}, o valor '
             'total estimado do pedido antes do embarque é de USD {{FOB_TOTAL}} ',
             False),
            ('({{FOB_TOTAL_EXTENSO}}).', True),
        ])

    # ── Sinal ────────────────────────────────────────────────────────────────────
    p = _achar(doc, lambda t: 'Sinal de 20%' in t)
    if p:
        _set_runs(p, [
            ('(i) Sinal de {{PCT_SINAL}}', True),
            (': no total de ', False),
            ('USD {{VALOR_SINAL_USD}} ({{VALOR_SINAL_USD_EXTENSO}}) ', True),
            ('utilizando a taxa do dólar fechado em ', False),
            ('R$ {{CAMBIO_SINAL}} ', True),
            ('no dia ', False),
            ('{{DATA_SINAL}}', True),
            (', totalizando o valor em reais de', False),
        ])

    # ── Valor em reais do sinal + vencimento + dados bancários ──────────────────
    p = _achar(doc, lambda t: 'que deverão ser pagos' in t)
    if p:
        _set_runs(p, [
            (' ', False),
            ('R$ {{VALOR_SINAL_BRL}} ({{VALOR_SINAL_BRL_EXTENSO}}) ', True),
            ('que deverão ser pagos, em uma única parcela, à vista até o dia ',
             False),
            ('{{DATA_VENC_SINAL}}', True),
            (' na conta da Pilar Imports: BANCO SANTANDER-033 / AGÊNCIA: 0108 / '
             'CONTA 13.008612-0 / CNPJ: 43.954.200/0001-96 - BANCO ITAU-341 / '
             'AGÊNCIA: 0196 / CONTA 99701-9 / CNPJ: 43.954.200/0001-96', False),
        ])

    # ── Saldo ────────────────────────────────────────────────────────────────────
    p = _achar(doc, lambda t: 'Saldo de 80%' in t)
    if p:
        _set_runs(p, [
            ('(ii) Saldo de {{PCT_SALDO}}', True),
            (': O valor integral da mercadoria, ou seja, o saldo, deverá ser pago '
             'e quitado por você em até no máximo {{DIAS_ANTES_DESEMBARQUE}} dias '
             'antes do desembarque do navio em território nacional, o que será '
             'devidamente comunicado pela Pilar Imports por e-mail e WhatsApp;',
             False),
        ])

    # ── Frete ────────────────────────────────────────────────────────────────────
    p = _achar(doc, lambda t: 'Frete Considerado previsto' in t)
    if p:
        _set_runs(p, [(
            'Frete Considerado previsto na transação é de USD ({{FRETE_USD}}) o '
            'container de 40hc, em caso de aumento ou redução será corrigido ao '
            'valor final;', False)])

    # ── Prazo de entrega ──────────────────────────────────────────────────────────
    p = _achar(doc, lambda t: 'Após o desembaraço da mercadoria' in t)
    if p:
        _set_runs(p, [
            ('Após o desembaraço da mercadoria, a ', False),
            ('Pilar Imports', True),
            (' se compromete a entregar a mercadoria na sede da sua empresa em até '
             '{{PRAZO_ENTREGA}} dias, contados do efetivo desembaraço da '
             'mercadoria. Caso isso não ocorra, a Pilar se compromete a efetuar '
             '100% da devolução do valor pago até aquele momento;', False),
        ])

    # ── Observações adicionais (parágrafo dedicado, antes da formalização) ──────
    p = _achar(doc, lambda t: 'Para formalização de entendimentos' in t)
    if p:
        _inserir_antes(p, [('{{OBSERVACOES}}', False)])

    doc.save(destino)

    # Diagnóstico: confere placeholders presentes
    chk = Document(destino)
    full = '\n'.join(p.text for p in chk.paragraphs)
    import re
    achados = sorted(set(re.findall(r'\{\{[A-Z_]+\}\}', full)))
    sys.stderr.write('Modelo salvo em: %s\n' % destino)
    sys.stderr.write('Placeholders encontrados (%d): %s\n'
                     % (len(achados), ', '.join(achados)))


if __name__ == '__main__':
    main()
